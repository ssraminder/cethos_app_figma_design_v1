#!/usr/bin/env python3
"""Minimal Markdown -> .docx renderer for Cethos controlled documents (staff
training guides, SOP working copies). NOT a full CommonMark parser - tuned to
our house markdown so the .docx tracks the .md single-source-of-truth.

Handles: `#`/`##`/`###` headings, `| pipe tables |` (with header shading and an
empty-header two-column "metadata" table), `- `/`* ` and `1. ` lists,
`**bold**`, `` `code` ``, fenced ``` code blocks (e.g. folder trees), markdown
links (rendered as their text), and `> `[SCREENSHOT N: ...]`` callouts (rendered
as shaded placeholder boxes, optionally annotated from a screenshot-map JSON).

Usage:
    python scripts/md_to_docx.py INPUT.md OUTPUT.docx [SCREENSHOT_MAP.json]

The screenshot map is {"<n>": "<capture note>"} keyed by the SCREENSHOT number;
mapped numbers show the capture note, unmapped ones show "[capture pending]".
Chrome-MCP captures live in the conversation transcript (see
docs/training/screenshots/INDEX.md) - the pixels are dropped into the boxes by
hand; this renderer lays out everything else.
"""
import sys
import re
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')
LINK = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
SS = re.compile(r'\[SCREENSHOT\s+(\d+)\s*:?\s*(.*?)\]', re.S)


def add_runs(p, text):
    """Add a paragraph's text, honouring **bold**, `code`, and [links](...)."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith('`'):
            r = p.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        else:
            lm = LINK.match(tok)
            r = p.add_run(lm.group(1) if lm else tok)
            r.font.color.rgb = RGBColor(0x0B, 0x5F, 0x6B)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def _shade(elem, fill):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill)
    elem.append(sh)


def shade_para(p, fill):
    _shade(p._p.get_or_add_pPr(), fill)


def shade_cell(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)


def split_row(line):
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]
    return [c.strip() for c in s.split('|')]


def is_sep(cells):
    return cells and all(re.fullmatch(r':?-{2,}:?', (c.strip() or '-')) for c in cells) \
        and any('-' in c for c in cells)


def render(src, out, ssmap):
    lines = open(src, encoding='utf-8').read().split('\n')
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()

        # fenced code block
        if s.startswith('```'):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph(); shade_para(p, 'F2F2F2')
            r = p.add_run('\n'.join(code)); r.font.name = 'Consolas'; r.font.size = Pt(9)
            continue

        # pipe table
        if s.startswith('|') and '|' in s[1:]:
            block = []
            while i < n and lines[i].strip().startswith('|'):
                block.append(lines[i]); i += 1
            rows = [split_row(b) for b in block]
            header, body = rows[0], rows[1:]
            if len(rows) >= 2 and is_sep(rows[1]):
                body = rows[2:]
            ncol = max(len(r) for r in rows)
            empty_header = all(c == '' for c in header)
            t = doc.add_table(rows=0, cols=ncol)
            try:
                t.style = 'Light Grid Accent 1'
            except KeyError:
                pass
            data = (body if empty_header else [header] + body)
            for ri, row in enumerate(data):
                cells = t.add_row().cells
                for ci in range(ncol):
                    cp = cells[ci].paragraphs[0]
                    add_runs(cp, row[ci] if ci < len(row) else '')
                    if not empty_header and ri == 0:
                        for rr in cp.runs:
                            rr.bold = True
                        shade_cell(cells[ci], 'D9E2F3')
                    elif empty_header and ci == 0:
                        for rr in cp.runs:
                            rr.bold = True
            doc.add_paragraph()
            continue

        if s.startswith('# '):
            doc.add_heading(s[2:].replace('**', ''), level=0)
        elif s.startswith('## '):
            doc.add_heading(s[3:].replace('**', ''), level=1)
        elif s.startswith('### '):
            doc.add_heading(s[4:].replace('**', ''), level=2)
        elif s == '---':
            pass
        elif s.startswith('>'):
            content = s.lstrip('> ').strip()
            mss = SS.search(content)
            if mss:
                num, desc = mss.group(1), mss.group(2).strip('` ')
                cap = ssmap.get(num)
                p = doc.add_paragraph(); shade_para(p, 'FFF2CC')
                p.add_run('[ SCREENSHOT %s ]  ' % num).bold = True
                add_runs(p, desc)
                tail = p.add_run('    %s' % (('capture: ' + cap) if cap else '[capture pending]'))
                tail.italic = True; tail.font.size = Pt(8.5)
            else:
                p = doc.add_paragraph(); add_runs(p, content)
                for r in p.runs:
                    r.italic = True
        elif re.match(r'^[-*] ', s):
            add_runs(doc.add_paragraph(style='List Bullet'), s[2:])
        elif re.match(r'^\d+\. ', s):
            add_runs(doc.add_paragraph(style='List Number'), re.sub(r'^\d+\. ', '', s))
        elif s == '':
            pass
        else:
            add_runs(doc.add_paragraph(), s)
        i += 1

    doc.save(out)
    print('wrote', out)


if __name__ == '__main__':
    src, out = sys.argv[1], sys.argv[2]
    ssmap = json.load(open(sys.argv[3], encoding='utf-8')) if len(sys.argv) > 3 else {}
    render(src, out, ssmap)
